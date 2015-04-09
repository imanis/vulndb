# -*- coding: utf-8 -*-
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Inches

from django.conf import settings as django_settings
from lxml.html import *
from vulnDB.projects.models import Project

from docx.enum.section import WD_ORIENTATION, WD_SECTION_START


def replace(document, toFind, toReplace):

    for i in document.paragraphs:
        tmpStr = i.text
        i.text = tmpStr.replace(toFind, toReplace)


def mergecontent(added_doc, document):

    for paragraph in added_doc.paragraphs:
        text = paragraph.text
        document.add_paragraph(text)



def formattext(text):
    #''.join(xml.etree.ElementTree.fromstring(text).itertext())
    return fromstring(text).text_content()


def generate_docx(projectid):
    savedir = django_settings.REPORTS_ROOT

    project = get_object_or_404(Project, pk=projectid)



    # document = Document()
    document = Document(''.join([savedir, 'template_rapport.docx']))

    # Generate Front Cover

    #replace(document,'_#nature#_', unicode(project.nature.nature))

    #replace(document,'_#nature#_', 'Technique')

        #document.add_heading(project.project, 9)


    # Generate Header

        #document.add_picture(django_settings.ROOT + project.client.get_picture(), width=Inches(1.25))

    #Generate footer


    #Generate Table of content

    h = document.add_heading(u'Table des matières', 1)
    h.style = "TBHeader"
    document.add_section(WD_SECTION_START.NEW_PAGE)


    h = document.add_heading(u'Table des tableaux', 1)
    h.style = "TBHeader"

    document.add_section(WD_SECTION_START.NEW_PAGE)

    h = document.add_heading(u'Table des figures', 1)
    h.style = "TBHeader"

    document.add_section(WD_SECTION_START.NEW_PAGE)




    # Generate synth Globale

    h = document.add_heading(u'Synthèse', 9)
    h.style = 'Titre1'


    h = document.add_heading(u'Synthèse Globale',4)
    h.style = 'Titre2'


    document.add_paragraph('Description de la situation ')

    document.add_paragraph('charts ')




    # Generate synth Detaillé

    tab_section = document.add_section(WD_SECTION_START.NEW_PAGE)


    h = document.add_heading(u'Synthèse détaillée',4)
    h.style = 'Titre2'

    table = document.add_table(rows=1, cols=9)
    table.style = 'AthenaTabStyle'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = unicode('Ref')
    hdr_cells[1].text = unicode('Elément impacté', 'utf-8')
    hdr_cells[2].text = unicode('Vulnérabilité', 'utf-8')
    hdr_cells[3].text = unicode('Niveau de criticité', 'utf-8')
    hdr_cells[4].text = unicode('Type', 'utf-8')
    hdr_cells[5].text = unicode('Conséquences d’exploitation', 'utf-8')
    hdr_cells[6].text = unicode('Recommandation', 'utf-8')
    hdr_cells[7].text = unicode('Priorité de mise en œuvre', 'utf-8')
    hdr_cells[8].text = unicode('Complexité de mise en œuvre', 'utf-8')

    for indx, v in enumerate(project.vulninst_set.all()):
        row_cells = table.add_row().cells
        row_cells[0].text = unicode(''.join(['V_',str(indx)]))
        row_cells[1].text = unicode(''.join(v.get_items_tags()))
        row_cells[2].text = unicode(v.vulnerability)
        row_cells[3].text = unicode(v.get_severity())
        row_cells[4].text = unicode(v.type)
        row_cells[5].text = unicode(v.exploitation_impact)
        row_cells[6].text = unicode(formattext(v.recommendation))
        row_cells[7].text = unicode(v.get_action_priority())
        row_cells[8].text = unicode(v.action_complexity)


    section = document.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE

    document.add_page_break()



    #Adding some Static parts
    #added_doc  = Document(''.join([savedir, 'demarche_test_intrusion.docx']))
    #mergecontent(added_doc, document)


    #Saving doc
    document.save(''.join([savedir, project.slug, '_', project.client.name, '.docx']))

    return document

